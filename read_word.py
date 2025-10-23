import docx
import sys

doc = docx.Document('Laporan Rancangan - Smart E-Arsip.docx')
for para in doc.paragraphs:
    print(para.text)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            print(cell.text)
